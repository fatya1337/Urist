import io
import logging
from pathlib import Path

from pypdf import PdfReader
from docx import Document

logger = logging.getLogger(__name__)


def read_pdf(source: bytes | str | Path) -> str:
    if isinstance(source, (str, Path)):
        reader = PdfReader(str(source))
    else:
        reader = PdfReader(io.BytesIO(source))

    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            pages.append(text)
        else:
            logger.warning("Страница %d: текст не извлечён", i + 1)

    full_text = "\n\n".join(pages).strip()
    if not full_text:
        raise ValueError("PDF не содержит извлекаемого текста (возможно, скан)")
    return full_text


def read_docx(source: bytes | str | Path) -> str:
    if isinstance(source, (str, Path)):
        doc = Document(str(source))
    else:
        doc = Document(io.BytesIO(source))

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs).strip()
    if not full_text:
        raise ValueError("DOCX не содержит текста")
    return full_text


def parse_document(filename: str, content: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return read_pdf(content)
    elif ext in (".docx", ".doc"):
        return read_docx(content)
    else:
        raise ValueError(f"Неподдерживаемый формат: {ext}. Принимаем PDF и DOCX.")
